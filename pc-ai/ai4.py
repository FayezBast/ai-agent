"""
JARVIS-Style AI Assistant - Enhanced & Secure Version
A comprehensive AI assistant with improved security, error handling, and features.
"""

import os
import sys
import json
import re
import subprocess
import platform
import webbrowser
import logging
import mimetypes
import difflib
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Union, Any
from functools import lru_cache
import time

# Third-party imports with error handling
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    print("⚠️ Google Generative AI not installed. Run: pip install google-generativeai")

try:
    import docx
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed. Run: pip install python-docx")

try:
    from openpyxl import Workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("⚠️ openpyxl not installed. Run: pip install openpyxl")

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️ fpdf2 not installed. Run: pip install fpdf2")

try:
    from rich.console import Console
    from rich.table import Table
    from rich.progress import Progress, SpinnerColumn, TextColumn
    from rich.panel import Panel
    RICH_AVAILABLE = True
    console = Console()
except ImportError:
    RICH_AVAILABLE = False
    print("⚠️ rich not installed. Run: pip install rich")


class Config:
    """Configuration settings for JARVIS AI"""
    WORKSPACE_DIR = Path.home() / "JARVIS_Workspace"
    LOG_FILE = "jarvis.log"
    HISTORY_FILE = "command_history.json"
    MAX_SEARCH_RESULTS = 10
    MAX_HISTORY_ENTRIES = 100
    
    SUPPORTED_FORMATS = ['docx', 'xlsx', 'pdf', 'txt', 'py', 'json', 'csv']
    
    APP_ALIASES = {
        "visual studio code": "code",
        "vscode": "code",
        "vs code": "code",
        "word": "winword",
        "excel": "excel",
        "notepad": "notepad",
        "calculator": "calc",
        "chrome": "chrome",
        "firefox": "firefox",
        "browser": "chrome"
    }
    
    VALID_INTENTS = [
        "file_creation",
        "file_management", 
        "system_control",
        "web_browse",
        "conversation",
        "help"
    ]


class JarvisLogger:
    """Enhanced logging for JARVIS"""
    
    def __init__(self, log_file: str = Config.LOG_FILE):
        self.setup_logging(log_file)
    
    def setup_logging(self, log_file: str):
        """Setup logging configuration"""
        log_format = '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        
        # Create logs directory if it doesn't exist
        log_path = Config.WORKSPACE_DIR / log_file
        log_path.parent.mkdir(exist_ok=True)
        
        logging.basicConfig(
            level=logging.INFO,
            format=log_format,
            handlers=[
                logging.FileHandler(log_path),
                logging.StreamHandler()
            ]
        )
        
        self.logger = logging.getLogger('JARVIS')
    
    def info(self, message: str):
        self.logger.info(message)
    
    def error(self, message: str):
        self.logger.error(message)
    
    def warning(self, message: str):
        self.logger.warning(message)


class SecurityValidator:
    """Security validation utilities"""
    
    @staticmethod
    def validate_filename(filename: str) -> str:
        """Validate filename for security and format"""
        if not filename or len(filename.strip()) == 0:
            raise ValueError("Filename cannot be empty")
        
        filename = filename.strip()
        
        # Check for invalid characters
        invalid_chars = '<>:"/\\|?*'
        if any(char in filename for char in invalid_chars):
            raise ValueError(f"Filename contains invalid characters: {invalid_chars}")
        
        # Check for reserved names (Windows)
        reserved_names = ['CON', 'PRN', 'AUX', 'NUL'] + [f'COM{i}' for i in range(1, 10)] + [f'LPT{i}' for i in range(1, 10)]
        if filename.upper() in reserved_names:
            raise ValueError(f"'{filename}' is a reserved filename")
        
        # Check length
        if len(filename) > 255:
            raise ValueError("Filename too long (max 255 characters)")
        
        return filename
    
    @staticmethod
    def validate_path(path: Union[str, Path]) -> Path:
        """Validate file path for security"""
        path = Path(path).resolve()
        
        # Ensure path is within allowed directories
        allowed_dirs = [Config.WORKSPACE_DIR, Path.home()]
        if not any(str(path).startswith(str(allowed_dir)) for allowed_dir in allowed_dirs):
            raise ValueError("Path is outside allowed directories")
        
        return path
    
    @staticmethod
    def sanitize_content(content: str) -> str:
        """Sanitize content for safe processing"""
        if not isinstance(content, str):
            content = str(content)
        
        # Remove potential script injections
        dangerous_patterns = [
            r'<script.*?</script>',
            r'javascript:',
            r'vbscript:',
            r'onload=',
            r'onerror='
        ]
        
        for pattern in dangerous_patterns:
            content = re.sub(pattern, '', content, flags=re.IGNORECASE)
        
        return content.strip()


class AIContentGenerator:
    """AI-powered content generation"""
    
    def __init__(self, api_key: Optional[str] = None):
        self.gemini_client = None
        self.logger = JarvisLogger().logger
        
        if api_key and GEMINI_AVAILABLE:
            self.setup_gemini(api_key)
    
    def setup_gemini(self, api_key: str):
        """Setup Google Gemini AI"""
        try:
            genai.configure(api_key=api_key)
            self.gemini_client = genai.GenerativeModel('gemini-1.5-flash-latest')
            self.logger.info("Google Gemini AI initialized successfully")
        except Exception as e:
            self.logger.error(f"Failed to initialize Gemini AI: {e}")
            self.gemini_client = None
    
    def generate_content(self, topic: str, output_format: str = 'text') -> str:
        """Generate content using AI or fallback templates"""
        if self.gemini_client:
            return self._generate_with_ai(topic, output_format)
        else:
            return self._generate_with_template(topic, output_format)
    
    def _generate_with_ai(self, topic: str, output_format: str) -> str:
        """Generate content using AI"""
        prompts = {
            'json': f"For '{topic}', generate structured data as a single JSON object with 'headers' (list) and 'rows' (list of lists). Make it realistic and useful.",
            'code': f"Write a complete, well-documented Python script for: '{topic}'. Include error handling and comments. IMPORTANT: Only output raw code, no markdown formatting.",
            'text': f"Write a comprehensive, well-structured document about '{topic}' using Markdown formatting (# headings, **bold**, * bullets, proper paragraphs)."
        }
        
        try:
            prompt = prompts.get(output_format, prompts['text'])
            self.logger.info(f"Generating {output_format} content for: {topic}")
            
            response = self.gemini_client.generate_content(prompt)
            return SecurityValidator.sanitize_content(response.text)
            
        except Exception as e:
            self.logger.error(f"AI content generation failed: {e}")
            return self._generate_with_template(topic, output_format)
    
    def _generate_with_template(self, topic: str, output_format: str) -> str:
        """Generate content using templates as fallback"""
        templates = {
            'text': f"""# {topic}

## Introduction
This document provides information about {topic}.

## Overview
{topic} is an important subject that requires detailed explanation.

## Key Points
* Point 1 about {topic}
* Point 2 about {topic}
* Point 3 about {topic}

## Conclusion
In conclusion, {topic} is a valuable topic for further study.
""",
            'code': f'''#!/usr/bin/env python3
"""
{topic} - Generated by JARVIS AI Assistant
"""

def main():
    """Main function for {topic}"""
    print("Hello from {topic}!")
    
    # TODO: Implement {topic} functionality
    pass

if __name__ == "__main__":
    main()
''',
            'json': json.dumps({
                "headers": ["Item", "Description", "Value"],
                "rows": [
                    [f"{topic} Item 1", "Description 1", "Value 1"],
                    [f"{topic} Item 2", "Description 2", "Value 2"],
                    [f"{topic} Item 3", "Description 3", "Value 3"]
                ]
            }, indent=2)
        }
        
        return templates.get(output_format, templates['text'])


class CommandAnalyzer:
    """Analyze user commands and extract intent"""
    
    def __init__(self, ai_generator: AIContentGenerator):
        self.ai_generator = ai_generator
        self.logger = JarvisLogger().logger
    
    @lru_cache(maxsize=50)
    def analyze_command(self, command: str) -> Dict[str, Any]:
        """Analyze command and return structured intent"""
        if self.ai_generator.gemini_client:
            return self._analyze_with_ai(command)
        else:
            return self._analyze_with_rules(command)
    
    def _analyze_with_ai(self, command: str) -> Dict[str, Any]:
        """Use AI to analyze command"""
        prompt = f"""
        You are a command analyzer. Convert the user command to a JSON object with this exact structure:

        COMMAND: "{command}"

        Valid intents: {', '.join(Config.VALID_INTENTS)}

        Action mapping rules:
        - "word doc" or ".docx" -> create_word
        - "excel" or ".xlsx" -> create_excel  
        - "pdf" -> create_pdf
        - "python" or ".py" -> create_python
        - "text file" or ".txt" -> create_text
        - "find" or "locate" -> find_file
        - "delete" or "remove" -> delete_file
        - "list files" -> list_files
        - "search" or "browse" -> web_search
        - "open" or "launch" -> open_application
        - "help" -> show_help
        - greetings -> chat

        Return only valid JSON:
        {{"intent": "...", "action": "...", "parameters": {{"filename": "...", "content": "...", "is_topic": true/false}}}}
        """
        
        try:
            response = self.ai_generator.gemini_client.generate_content(prompt)
            match = re.search(r'\{.*\}', response.text, re.DOTALL)
            if match:
                result = json.loads(match.group(0))
                return self._validate_analysis(result)
        except Exception as e:
            self.logger.error(f"AI command analysis failed: {e}")
        
        return self._analyze_with_rules(command)
    
    def _analyze_with_rules(self, command: str) -> Dict[str, Any]:
        """Fallback rule-based command analysis"""
        cmd = command.lower().strip()
        
        # File creation patterns
        if any(word in cmd for word in ['create', 'make', 'generate', 'write']):
            if any(word in cmd for word in ['word', 'doc', '.docx']):
                return self._create_analysis('file_creation', 'create_word', command)
            elif any(word in cmd for word in ['excel', '.xlsx', 'spreadsheet']):
                return self._create_analysis('file_creation', 'create_excel', command)
            elif 'pdf' in cmd:
                return self._create_analysis('file_creation', 'create_pdf', command)
            elif any(word in cmd for word in ['python', '.py', 'script']):
                return self._create_analysis('file_creation', 'create_python', command)
            elif any(word in cmd for word in ['text', '.txt', 'file']):
                return self._create_analysis('file_creation', 'create_text', command)
        
        # File management
        elif any(word in cmd for word in ['find', 'locate', 'search for']):
            return {"intent": "file_management", "action": "find_file", "parameters": {"query": cmd}}
        elif any(word in cmd for word in ['delete', 'remove']):
            return {"intent": "file_management", "action": "delete_file", "parameters": {"query": cmd}}
        elif 'list files' in cmd:
            return {"intent": "file_management", "action": "list_files", "parameters": {}}
        
        # System control
        elif any(word in cmd for word in ['open', 'launch', 'start']):
            app = re.sub(r'\b(open|launch|start)\b', '', cmd).strip()
            return {"intent": "system_control", "action": "open_application", "parameters": {"application": app}}
        
        # Web browsing
        elif any(word in cmd for word in ['search', 'browse', 'google']):
            query = re.sub(r'\b(search|browse|google)\b', '', cmd).strip()
            return {"intent": "web_browse", "action": "web_search", "parameters": {"search_query": query}}
        
        # Help
        elif any(word in cmd for word in ['help', 'commands', 'what can you do']):
            return {"intent": "help", "action": "show_help", "parameters": {}}
        
        # Conversation
        else:
            return {"intent": "conversation", "action": "chat", "parameters": {"message": command}}
    
    def _create_analysis(self, intent: str, action: str, command: str) -> Dict[str, Any]:
        """Create file creation analysis"""
        # Extract topic/content
        topic_patterns = [
            r'(?:about|on|for|regarding)\s+(.+)',
            r'(?:create|make|write)\s+(?:a\s+)?(?:word|excel|pdf|python|text)?\s*(?:file|doc|document|script)?\s*(?:about|on|for|regarding)?\s*(.+)',
        ]
        
        topic = None
        for pattern in topic_patterns:
            match = re.search(pattern, command, re.IGNORECASE)
            if match:
                topic = match.group(1).strip()
                break
        
        if not topic:
            topic = command
        
        # Generate filename
        filename = self._generate_filename(topic, action)
        
        return {
            "intent": intent,
            "action": action,
            "parameters": {
                "filename": filename,
                "content": topic,
                "is_topic": True
            }
        }
    
    def _generate_filename(self, topic: str, action: str) -> str:
        """Generate appropriate filename from topic"""
        # Clean topic for filename
        clean_topic = re.sub(r'[^\w\s-]', '', topic)
        clean_topic = re.sub(r'\s+', '_', clean_topic)
        clean_topic = clean_topic[:50]  # Limit length
        
        extensions = {
            'create_word': '.docx',
            'create_excel': '.xlsx', 
            'create_pdf': '.pdf',
            'create_python': '.py',
            'create_text': '.txt'
        }
        
        ext = extensions.get(action, '.txt')
        return f"{clean_topic}{ext}"
    
    def _validate_analysis(self, analysis: Dict[str, Any]) -> Dict[str, Any]:
        """Validate analysis result"""
        if analysis.get('intent') not in Config.VALID_INTENTS:
            analysis['intent'] = 'conversation'
            analysis['action'] = 'chat'
        
        return analysis


class FileManager:
    """Handle file operations"""
    
    def __init__(self, workspace_dir: Path):
        self.workspace_dir = workspace_dir
        self.workspace_dir.mkdir(exist_ok=True)
        self.logger = JarvisLogger().logger
    
    def create_file(self, filepath: Path, content: str, file_type: str) -> str:
        """Create file with appropriate format"""
        try:
            filepath = SecurityValidator.validate_path(filepath)
            content = SecurityValidator.sanitize_content(content)
            
            if file_type == 'docx' and DOCX_AVAILABLE:
                return self._create_word_document(filepath, content)
            elif file_type == 'xlsx' and EXCEL_AVAILABLE:
                return self._create_excel_file(filepath, content)
            elif file_type == 'pdf' and PDF_AVAILABLE:
                return self._create_pdf_file(filepath, content)
            else:
                return self._create_text_file(filepath, content)
                
        except Exception as e:
            self.logger.error(f"File creation failed: {e}")
            raise
    
    def _create_text_file(self, filepath: Path, content: str) -> str:
        """Create text file"""
        filepath.write_text(content, encoding='utf-8')
        return f"Text file created: {filepath.name}"
    
    def _create_word_document(self, filepath: Path, content: str) -> str:
        """Create formatted Word document"""
        doc = docx.Document()
        
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('* ') or line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)
        
        doc.save(filepath)
        return f"Word document created: {filepath.name}"
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add formatted text to Word paragraph"""
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                paragraph.add_run(part[2:-2]).bold = True
            elif part.startswith('*') and part.endswith('*'):
                paragraph.add_run(part[1:-1]).italic = True
            else:
                paragraph.add_run(part)
    
    def _create_excel_file(self, filepath: Path, content: str) -> str:
        """Create Excel file from JSON data"""
        try:
            # Try to parse as JSON first
            data = json.loads(content)
            if 'headers' in data and 'rows' in data:
                workbook = Workbook()
                sheet = workbook.active
                sheet.append(data['headers'])
                for row in data['rows']:
                    sheet.append(row)
                workbook.save(filepath)
                return f"Excel file created: {filepath.name}"
        except json.JSONDecodeError:
            pass
        
        # Fallback: create simple Excel with content
        workbook = Workbook()
        sheet = workbook.active
        sheet['A1'] = 'Content'
        sheet['A2'] = content
        workbook.save(filepath)
        return f"Excel file created: {filepath.name}"
    
    def _create_pdf_file(self, filepath: Path, content: str) -> str:
        """Create PDF file"""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Try to use Unicode font
        try:
            pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
            pdf.set_font('DejaVu', '', 12)
        except:
            pdf.set_font('Arial', '', 12)
        
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                pdf.ln(5)
                continue
                
            if line.startswith('# '):
                pdf.set_font(size=18)
                pdf.multi_cell(0, 12, line[2:], new_x="LMARGIN", new_y="NEXT")
                pdf.set_font(size=12)
            elif line.startswith('## '):
                pdf.set_font(size=14)
                pdf.multi_cell(0, 10, line[3:], new_x="LMARGIN", new_y="NEXT")
                pdf.set_font(size=12)
            elif line.startswith('* ') or line.startswith('- '):
                pdf.multi_cell(0, 8, f'• {line[2:]}', new_x="LMARGIN", new_y="NEXT")
            else:
                pdf.multi_cell(0, 8, line, new_x="LMARGIN", new_y="NEXT")
        
        pdf.output(filepath)
        return f"PDF file created: {filepath.name}"
    
    def find_files(self, query: str, max_results: int = Config.MAX_SEARCH_RESULTS) -> List[Path]:
        """Find files matching query"""
        results = []
        search_dirs = [self.workspace_dir, Path.home() / "Documents", Path.home() / "Desktop"]
        
        for search_dir in search_dirs:
            if search_dir.exists():
                try:
                    matches = list(search_dir.rglob(f"*{query}*"))
                    results.extend(matches[:max_results])
                    if len(results) >= max_results:
                        break
                except PermissionError:
                    continue
        
        return results[:max_results]
    
    def delete_file(self, filepath: Path) -> str:
        """Delete file safely"""
        try:
            filepath = SecurityValidator.validate_path(filepath)
            if filepath.exists():
                filepath.unlink()
                return f"File deleted: {filepath.name}"
            else:
                return "File not found"
        except Exception as e:
            self.logger.error(f"File deletion failed: {e}")
            raise
    
    def list_workspace_files(self) -> List[Path]:
        """List files in workspace"""
        try:
            return [f for f in self.workspace_dir.iterdir() if f.is_file()]
        except Exception as e:
            self.logger.error(f"Failed to list files: {e}")
            return []
    
    def open_file(self, filepath: Path) -> str:
        """Open file with system default application"""
        try:
            system = platform.system().lower()
            if system == "windows":
                os.startfile(filepath)
            elif system == "darwin":
                subprocess.run(["open", filepath], check=True)
            else:
                subprocess.run(["xdg-open", filepath], check=True)
            return "File opened successfully"
        except Exception as e:
            self.logger.error(f"Failed to open file: {e}")
            return f"Could not open file: {e}"


class SystemController:
    """Handle system operations"""
    
    def __init__(self):
        self.logger = JarvisLogger().logger
    
    def open_application(self, app_name: str) -> str:
        """Open application by name"""
        try:
            # Get actual command from aliases
            command = Config.APP_ALIASES.get(app_name.lower(), app_name)
            
            system = platform.system().lower()
            if system == "windows":
                os.system(f'start {command}')
            else:
                subprocess.Popen([command], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            
            return f"Opened {app_name}"
            
        except Exception as e:
            self.logger.error(f"Failed to open {app_name}: {e}")
            return f"Could not open {app_name}: {e}"


class WebController:
    """Handle web operations"""
    
    def __init__(self):
        self.logger = JarvisLogger().logger
    
    def search_web(self, query: str) -> str:
        """Search the web"""
        try:
            if query:
                search_url = f"https://www.google.com/search?q={query.replace(' ', '+')}"
                webbrowser.open(search_url)
                return f"Searching the web for: {query}"
            else:
                webbrowser.open("https://www.google.com")
                return "Opened Google homepage"
        except Exception as e:
            self.logger.error(f"Web search failed: {e}")
            return f"Could not perform web search: {e}"


class CommandHistory:
    """Manage command history"""
    
    def __init__(self, history_file: Path):
        self.history_file = history_file
        self.history = []
        self.load_history()
    
    def load_history(self):
        """Load command history from file"""
        try:
            if self.history_file.exists():
                with open(self.history_file, 'r') as f:
                    self.history = json.load(f)
        except Exception:
            self.history = []
    
    def save_history(self):
        """Save command history to file"""
        try:
            # Keep only recent entries
            self.history = self.history[-Config.MAX_HISTORY_ENTRIES:]
            with open(self.history_file, 'w') as f:
                json.dump(self.history, f, indent=2)
        except Exception:
            pass
    
    def add_command(self, command: str, success: bool, result: str = ""):
        """Add command to history"""
        entry = {
            'timestamp': datetime.now().isoformat(),
            'command': command,
            'success': success,
            'result': result[:100] if result else ""  # Truncate long results
        }
        self.history.append(entry)
        self.save_history()


class JarvisAI:
    """Main JARVIS AI Assistant class"""
    
    def __init__(self, api_key: Optional[str] = None):
        self.logger = JarvisLogger()
        self.system = platform.system().lower()
        self.workspace_dir = Config.WORKSPACE_DIR
        self.workspace_dir.mkdir(exist_ok=True)
        
        # Initialize components
        self.ai_generator = AIContentGenerator(api_key)
        self.command_analyzer = CommandAnalyzer(self.ai_generator)
        self.file_manager = FileManager(self.workspace_dir)
        self.system_controller = SystemController()
        self.web_controller = WebController()
        self.command_history = CommandHistory(self.workspace_dir / Config.HISTORY_FILE)
        
        self.pending_confirmation = None
        
        self._display_startup_info()
    
    def _display_startup_info(self):
        """Display startup information"""
        if RICH_AVAILABLE:
            startup_panel = Panel(
                f"🤖 JARVIS AI Assistant - Enhanced Version\n"
                f"💻 System: {platform.system()}\n"
                f"🏠 Workspace: {self.workspace_dir}\n"
                f"🧠 AI: {'✅ Enabled' if self.ai_generator.gemini_client else '❌ Disabled'}\n"
                f"📊 Features: {len([x for x in [DOCX_AVAILABLE, EXCEL_AVAILABLE, PDF_AVAILABLE, RICH_AVAILABLE] if x])}/4 modules loaded",
                title="JARVIS Initialized",
                style="bold green"
            )
            console.print(startup_panel)
        else:
            print("🤖 JARVIS AI Assistant - Enhanced Version")
            print(f"💻 System: {platform.system()}")
            print(f"🏠 Workspace: {self.workspace_dir}")
            print(f"🧠 AI: {'✅ Enabled' if self.ai_generator.gemini_client else '❌ Disabled'}")
    
    def suggest_commands(self, invalid_command: str) -> str:
        """Suggest similar valid commands"""
        common_commands = [
            'create word document', 'create excel file', 'create pdf',
            'create python script', 'find file', 'list files',
            'open application', 'search web', 'help'
        ]
        
        suggestions = difflib.get_close_matches(invalid_command, common_commands, n=3, cutoff=0.6)
        if suggestions:
            return f"Did you mean: {', '.join(suggestions)}?"
        return "Type 'help' for available commands."
    
    def execute_command(self, analysis: Dict[str, Any]) -> str:
        """Execute command based on analysis"""
        intent = analysis.get("intent", "unknown")
        action = analysis.get("action", "unknown")
        params = analysis.get("parameters", {})
        
        try:
            if intent == "file_creation":
                return self._handle_file_creation(action, params)
            elif intent == "file_management":
                return self._handle_file_management(action, params)
            elif intent == "system_control":
                return self._handle_system_control(action, params)
            elif intent == "web_browse":
                return self._handle_web_browse(action, params)
            elif intent == "conversation":
                return self._handle_conversation(params)
            elif intent == "help":
                return self._show_help()
            else:
                return self.suggest_commands(params.get('message', ''))
                
        except Exception as e:
            self.logger.error(f"Command execution failed: {e}")
            return f"An error occurred: {e}"
    
    def _handle_file_creation(self, action: str, params: Dict[str, Any]) -> str:
        """Handle file creation commands"""
        filename = params.get("filename")
        content = params.get("content", "")
        is_topic = params.get("is_topic", False)
        
        if not filename:
            return "❌ Filename is required for file creation"
        
        try:
            # Validate filename
            filename = SecurityValidator.validate_filename(filename)
            filepath = self.workspace_dir / filename
            
            # Generate content if needed
            if is_topic and content:
                if RICH_AVAILABLE:
                    with console.status(f"[bold green]Generating content for '{content}'..."):
                        time.sleep(0.5)  # Brief pause for visual effect
                        if action == "create_excel":
                            generated_content = self.ai_generator.generate_content(content, 'json')
                        elif action == "create_python":
                            generated_content = self.ai_generator.generate_content(content, 'code')
                        else:
                            generated_content = self.ai_generator.generate_content(content, 'text')
                else:
                    print(f"🧠 Generating content for '{content}'...")
                    if action == "create_excel":
                        generated_content = self.ai_generator.generate_content(content, 'json')
                    elif action == "create_python":
                        generated_content = self.ai_generator.generate_content(content, 'code')
                    else:
                        generated_content = self.ai_generator.generate_content(content, 'text')
            else:
                generated_content = content
            
            # Determine file type
            file_type = {
                'create_word': 'docx',
                'create_excel': 'xlsx',
                'create_pdf': 'pdf',
                'create_python': 'py',
                'create_text': 'txt'
            }.get(action, 'txt')
            
            # Create the file
            result = self.file_manager.create_file(filepath, generated_content, file_type)
            
            # Try to open the file
            open_result = self.file_manager.open_file(filepath)
            if "successfully" in open_result:
                result += " and opened it for you."
            else:
                result += f" (Note: {open_result})"
            
            return f"✅ {result}"
            
        except ValueError as e:
            return f"❌ Invalid input: {e}"
        except Exception as e:
            self.logger.error(f"File creation error: {e}")
            return f"❌ Failed to create file: {e}"
    
    def _handle_file_management(self, action: str, params: Dict[str, Any]) -> str:
        """Handle file management commands"""
        if action == "list_files":
            files = self.file_manager.list_workspace_files()
            if not files:
                return "📁 Your workspace is empty."
            
            if RICH_AVAILABLE:
                table = Table(title="Workspace Files")
                table.add_column("Name", style="cyan")
                table.add_column("Size", style="green")
                table.add_column("Modified", style="yellow")
                
                for file in files:
                    try:
                        stat = file.stat()
                        size = f"{stat.st_size:,} bytes"
                        modified = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M")
                        table.add_row(file.name, size, modified)
                    except:
                        table.add_row(file.name, "Unknown", "Unknown")
                
                console.print(table)
                return ""
            else:
                result = "📁 Files in workspace:\n"
                for file in files:
                    try:
                        size = file.stat().st_size
                        result += f"  📄 {file.name} ({size:,} bytes)\n"
                    except:
                        result += f"  📄 {file.name}\n"
                return result
        
        elif action == "find_file":
            query = params.get("query", "").replace("find ", "").replace("locate ", "").strip()
            if not query:
                return "❌ Please specify what file to find."
            
            if RICH_AVAILABLE:
                with console.status(f"[bold yellow]Searching for '{query}'..."):
                    results = self.file_manager.find_files(query)
            else:
                print(f"🔍 Searching for '{query}'...")
                results = self.file_manager.find_files(query)
            
            if not results:
                return f"❌ No files found matching '{query}'"
            
            result = f"🔍 Found {len(results)} file(s) matching '{query}':\n"
            for file_path in results:
                result += f"  📄 {file_path}\n"
            
            return result
        
        elif action == "delete_file":
            query = params.get("query", "").replace("delete ", "").replace("remove ", "").strip()
            if not query:
                return "❌ Please specify what file to delete."
            
            results = self.file_manager.find_files(query)
            if not results:
                return f"❌ No files found matching '{query}'"
            
            if len(results) > 1:
                result = f"⚠️ Found multiple files matching '{query}':\n"
                for i, file_path in enumerate(results[:5], 1):
                    result += f"  {i}. {file_path}\n"
                result += "\nPlease be more specific or use the full filename."
                return result
            
            # Single file found - ask for confirmation
            file_to_delete = results[0]
            self.pending_confirmation = {
                "action": "confirm_delete",
                "file_path": file_to_delete,
                "prompt": f"⚠️ Are you sure you want to permanently delete '{file_to_delete.name}'? (yes/no)"
            }
            return self.pending_confirmation["prompt"]
        
        return "❌ Unknown file management action."
    
    def _handle_system_control(self, action: str, params: Dict[str, Any]) -> str:
        """Handle system control commands"""
        if action == "open_application":
            app_name = params.get("application", "").strip()
            if not app_name:
                return "❌ Please specify an application to open."
            
            return f"🚀 {self.system_controller.open_application(app_name)}"
        
        return "❌ Unknown system control action."
    
    def _handle_web_browse(self, action: str, params: Dict[str, Any]) -> str:
        """Handle web browsing commands"""
        if action == "web_search":
            query = params.get("search_query", "").strip()
            return f"🌐 {self.web_controller.search_web(query)}"
        
        return "❌ Unknown web browsing action."
    
    def _handle_conversation(self, params: Dict[str, Any]) -> str:
        """Handle conversational commands"""
        message = params.get("message", "").lower()
        
        responses = {
            "hi": "Hello! I'm JARVIS, your AI assistant. How can I help you today?",
            "hello": "Hello! I'm JARVIS, your AI assistant. How can I help you today?",
            "how are you": "I am functioning at optimal capacity and ready to assist you.",
            "what can you do": "I can help you create documents, manage files, open applications, search the web, and much more. Type 'help' for detailed commands.",
            "thanks": "You're welcome! Always happy to help.",
            "thank you": "You're welcome! Always happy to help.",
            "bye": "Goodbye! It was a pleasure assisting you.",
            "goodbye": "Goodbye! It was a pleasure assisting you."
        }
        
        for key, response in responses.items():
            if key in message:
                return f"🤖 {response}"
        
        # Default conversation response
        if self.ai_generator.gemini_client:
            try:
                response = self.ai_generator.gemini_client.generate_content(
                    f"You are JARVIS, an AI assistant. Respond to: '{message}' in a helpful, professional manner."
                )
                return f"🤖 {response.text}"
            except:
                pass
        
        return "🤖 I'm JARVIS, your AI assistant. I can help you create files, manage documents, open applications, and more. Type 'help' to see what I can do!"
    
    def _show_help(self) -> str:
        """Show help information"""
        help_text = """
🤖 JARVIS AI Assistant - Available Commands

📝 FILE CREATION:
  • "create a word document about [topic]"
  • "create an excel file about [topic]"  
  • "create a pdf report on [topic]"
  • "create a python script for [purpose]"
  • "create a text file about [topic]"

📁 FILE MANAGEMENT:
  • "list files" - Show workspace files
  • "find [filename]" - Search for files
  • "delete [filename]" - Delete a file

🖥️ SYSTEM CONTROL:
  • "open [application]" - Launch applications
  • Available apps: VS Code, Word, Excel, Chrome, Firefox, etc.

🌐 WEB BROWSING:
  • "search [query]" - Search the web
  • "browse [topic]" - Open web search

💬 CONVERSATION:
  • Ask questions, get help, or just chat!

🔧 SPECIAL FEATURES:
  • Command history tracking
  • Rich visual output (when available)
  • AI-powered content generation
  • Security validation
  • Error recovery

Type any command naturally - I'll understand what you want to do!
        """
        
        if RICH_AVAILABLE:
            help_panel = Panel(help_text, title="JARVIS Help", style="bold blue")
            console.print(help_panel)
            return ""
        else:
            return help_text
    
    def _confirm_delete_action(self, file_path: Path) -> str:
        """Confirm and execute file deletion"""
        try:
            result = self.file_manager.delete_file(file_path)
            return f"✅ {result}"
        except Exception as e:
            return f"❌ Failed to delete file: {e}"
    
    def run(self):
        """Main application loop"""
        if RICH_AVAILABLE:
            console.print("\n[bold green]JARVIS AI Assistant is active. How may I assist you?[/bold green]")
        else:
            print("\n🤖 JARVIS AI Assistant is active. How may I assist you?")
        
        while True:
            try:
                if RICH_AVAILABLE:
                    console.print("\n" + "="*60)
                else:
                    print("\n" + "="*50)
                
                # Handle pending confirmations
                if self.pending_confirmation:
                    if RICH_AVAILABLE:
                        user_response = console.input("[yellow]Your confirmation (yes/no): [/yellow]").strip().lower()
                    else:
                        user_response = input("💬 Your confirmation (yes/no): ").strip().lower()
                    
                    if user_response in ['yes', 'y']:
                        if self.pending_confirmation["action"] == "confirm_delete":
                            result = self._confirm_delete_action(self.pending_confirmation["file_path"])
                            if RICH_AVAILABLE:
                                console.print(f"[bold green]🤖 JARVIS:[/bold green] {result}")
                            else:
                                print(f"🤖 JARVIS: {result}")
                    else:
                        if RICH_AVAILABLE:
                            console.print("[bold yellow]🤖 JARVIS:[/bold yellow] Action cancelled.")
                        else:
                            print("🤖 JARVIS: Action cancelled.")
                    
                    self.pending_confirmation = None
                    continue
                
                # Get user command
                if RICH_AVAILABLE:
                    command = console.input("[cyan]💬 Type your command (or 'exit' to quit): [/cyan]").strip()
                else:
                    command = input("💬 Type your command (or 'exit' to quit): ").strip()
                
                if not command:
                    continue
                
                if command.lower() in ['exit', 'quit', 'shutdown', 'bye']:
                    break
                
                # Analyze and execute command
                if RICH_AVAILABLE:
                    with console.status("[bold green]Processing command..."):
                        analysis = self.command_analyzer.analyze_command(command)
                else:
                    print("🧠 Processing command...")
                    analysis = self.command_analyzer.analyze_command(command)
                
                # Show analysis (optional debug info)
                intent = analysis.get('intent', 'N/A')
                action = analysis.get('action', 'N/A')
                
                if RICH_AVAILABLE:
                    console.print(f"[dim]📊 Intent: {intent} | Action: {action}[/dim]")
                else:
                    print(f"📊 Intent: {intent} | Action: {action}")
                
                # Execute command
                result = self.execute_command(analysis)
                success = not result.startswith("❌")
                
                # Add to history
                self.command_history.add_command(command, success, result)
                
                # Display result
                if self.pending_confirmation:
                    # Result contains confirmation prompt
                    if RICH_AVAILABLE:
                        console.print(f"[bold red]🤖 JARVIS:[/bold red] {result}")
                    else:
                        print(f"🤖 JARVIS: {result}")
                else:
                    # Regular result
                    if result:  # Only print if there's content
                        if RICH_AVAILABLE:
                            style = "bold green" if success else "bold red"
                            console.print(f"[{style}]🤖 JARVIS:[/{style}] {result}")
                        else:
                            print(f"🤖 JARVIS: {result}")
                
            except KeyboardInterrupt:
                if RICH_AVAILABLE:
                    console.print("\n[yellow]Interrupted by user[/yellow]")
                else:
                    print("\nInterrupted by user")
                break
            except Exception as e:
                self.logger.error(f"Unexpected error: {e}")
                if RICH_AVAILABLE:
                    console.print(f"[bold red]❌ An unexpected error occurred: {e}[/bold red]")
                else:
                    print(f"❌ An unexpected error occurred: {e}")
        
        # Shutdown message
        if RICH_AVAILABLE:
            shutdown_panel = Panel(
                "Thank you for using JARVIS AI Assistant!\nGoodbye! 👋",
                title="Shutting Down",
                style="bold blue"
            )
            console.print(shutdown_panel)
        else:
            print("\n🤖 Shutting down JARVIS. Goodbye! 👋")


def main():
    """Main entry point"""
    if RICH_AVAILABLE:
        console.print("[bold cyan]🚀 Initializing JARVIS AI Assistant...[/bold cyan]")
    else:
        print("🚀 Initializing JARVIS AI Assistant...")
    
    # Check for required packages
    missing_packages = []
    if not GEMINI_AVAILABLE:
        missing_packages.append("google-generativeai")
    if not DOCX_AVAILABLE:
        missing_packages.append("python-docx")
    if not EXCEL_AVAILABLE:
        missing_packages.append("openpyxl")
    if not PDF_AVAILABLE:
        missing_packages.append("fpdf2")
    if not RICH_AVAILABLE:
        missing_packages.append("rich")
    
    if missing_packages:
        print(f"📋 Optional packages not installed: {', '.join(missing_packages)}")
        print(f"💡 Install with: pip install {' '.join(missing_packages)}")
        print("⚠️ Some features may be limited without these packages.\n")
    
    # Get API key from environment
    api_key = os.getenv('GOOGLE_API_KEY')
    
    if not api_key:
        if RICH_AVAILABLE:
            console.print("[yellow]🔑 WARNING: GOOGLE_API_KEY environment variable not set![/yellow]")
            console.print("[dim]AI features will be limited. Set the environment variable to enable full AI capabilities.[/dim]\n")
        else:
            print("🔑 WARNING: GOOGLE_API_KEY environment variable not set!")
            print("AI features will be limited. Set the environment variable to enable full AI capabilities.\n")
    
    try:
        # Initialize and run JARVIS
        jarvis = JarvisAI(api_key=api_key)
        jarvis.run()
        
    except Exception as e:
        if RICH_AVAILABLE:
            console.print(f"[bold red]❌ Failed to start JARVIS: {e}[/bold red]")
        else:
            print(f"❌ Failed to start JARVIS: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()