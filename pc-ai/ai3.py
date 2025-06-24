"""
JARVIS-Style AI Assistant with Natural Language Processing (Upgraded v7 - The 'Intent Lock' Update)
Forces the AI to use a strict, predefined list of intents for maximum reliability.
"""

import os
import sys
import json
import re
import subprocess
import platform
import webbrowser
from pathlib import Path
from datetime import datetime
import google.generativeai as genai
import docx
from docx.shared import Inches
from openpyxl import Workbook
from fpdf import FPDF

class JarvisAI:
    def __init__(self, api_key):
        self.system = platform.system().lower()
        self.workspace_dir = Path.home() / "JARVIS_Workspace"
        self.workspace_dir.mkdir(exist_ok=True)
        
        self.gemini_client = None
        self.pending_confirmation = None
        
        self.setup_ai_models(api_key)
        print(f"🤖 JARVIS AI Assistant initialized")
        print(f"💻 System: {platform.system()}")
        print(f"🏠 Workspace: {self.workspace_dir}")
    
    def setup_ai_models(self, api_key):
        try:
            if api_key and api_key != "PASTE_YOUR_REAL_API_KEY_HERE":
                genai.configure(api_key=api_key)
                self.gemini_client = genai.GenerativeModel('gemini-1.5-flash-latest')
                print("✅ Google Gemini API connected")
            else:
                print("⚠️ Google Gemini API key not provided. AI features will be limited.")
        except Exception as e:
            print(f"⚠️ AI models setup error: {e}")

    def _generate_content_with_ai(self, topic, output_format='text'):
        if not self.gemini_client: return "Cannot generate content without a valid API key."
        if output_format == 'json':
            prompt = f"For '{topic}', generate structured data as a single JSON object with 'headers' (list) and 'rows' (list of lists)."
        elif output_format == 'code':
            prompt = f"Write a complete, executable script for: '{topic}'. IMPORTANT: Only output raw code, no markdown."
        else:
            prompt = f"Write a well-structured document about '{topic}' using Markdown (# headings, **bold**, * bullets)."
        try:
            print(f"🧠 Generating content for '{topic}' in {output_format} format...")
            response = self.gemini_client.generate_content(prompt)
            return response.text
        except Exception as e:
            return f"An error occurred during content generation: {e}"

    ### --- NEW, ULTRA-STRICT ANALYSIS PROMPT --- ###
    def analyze_command_with_ai(self, command):
        if not self.gemini_client: return self._analyze_with_rules(command)

        prompt = f"""
        You are a command-to-JSON converter. Your only job is to analyze the user's command and return a single, valid JSON object based on the strict rules below.

        COMMAND: "{command}"

        **RULE 1: The 'intent' MUST be one of these exact strings:**
        - "file_creation"
        - "file_management"
        - "system_control"
        - "web_Browse"
        - "conversation"

        **RULE 2: The 'action' MUST be chosen based on these keywords:**
        - "word doc" or ".docx" -> `create_word`
        - "excel file" or ".xlsx" -> `create_excel`
        - "pdf" -> `create_pdf`
        - "python script" or ".py" -> `create_python`
        - "text file" or ".txt" -> `create_text`
        - "find" or "locate" -> `find_file`
        - "delete" or "remove" -> `delete_file`
        - "search" or "browse" -> `web_search`
        - "open" or "launch" -> `open_application`
        - "hi", "hello", "how are you", "joke" -> `chat`

        **RULE 3: Parameter Extraction:**
        - If the user asks to write **about** a subject, `content` must be the subject itself, and `is_topic` must be `true`.
        - `search_query` must contain the full search phrase.
        - If no filename is given, create a logical one from the topic.

        **JSON STRUCTURE:**
        {{"intent": "...", "action": "...", "parameters": {{"filename": "...", "content": "...", "is_topic": true/false, ...}}}}

        **Examples:**
        - Command: "create a pdf report on the history of AI"
          -> {{"intent": "file_creation", "action": "create_pdf", "parameters": {{"filename": "history_of_AI.pdf", "content": "the history of AI", "is_topic": true}}}}
        - Command: "create an excel file of the top 5 cryptocurrencies"
          -> {{"intent": "file_creation", "action": "create_excel", "parameters": {{"filename": "top_5_cryptocurrencies.xlsx", "content": "top 5 cryptocurrencies", "is_topic": true}}}}
        - Command: "hi how are you"
          -> {{"intent": "conversation", "action": "chat", "parameters": {{"content": "hi how are you", "is_topic": false}}}}
        """
        try:
            print("🧠 Analyzing with Gemini...")
            response = self.gemini_client.generate_content(prompt)
            # Use regex to find the JSON block robustly
            match = re.search(r'\{.*\}', response.text, re.DOTALL)
            if not match: raise ValueError("No valid JSON object found in AI response.")
            json_response_string = match.group(0)
            result = json.loads(json_response_string)
            return result
        except Exception as e:
            print(f"Gemini analysis error: {e}")
            return self._analyze_with_rules(command)

    def _analyze_with_rules(self, command):
        cmd = command.lower()
        if "open" in cmd or "launch" in cmd:
            app = cmd.replace("open", "").replace("launch","").strip()
            return {"intent": "system_control", "action": "open_application", "parameters": {"application": app}}
        return {"intent": "conversation", "action": "chat", "parameters": {"message": command}}

    def execute_command(self, analysis):
        intent = analysis.get("intent", "unknown")
        action = analysis.get("action", "unknown")
        params = analysis.get("parameters", {})
        try:
            if intent == "file_creation": return self._handle_file_creation(action, params)
            elif intent == "file_management": return self._handle_file_management(action, params)
            elif intent == "system_control": return self._handle_system_control(action, params)
            elif intent == "web_Browse": return self._handle_web_Browse(action, params)
            elif intent == "conversation": return self._handle_conversation(params)
            else:
                # This block should no longer be reached if the prompt works correctly
                return "I'm not sure how to help with that."
        except Exception as e:
            return f"Error executing command: {e}"

    def _open_file(self, filepath):
        try:
            if self.system == "windows": os.startfile(filepath)
            elif self.system == "darwin": subprocess.run(["open", filepath], check=True)
            else: subprocess.run(["xdg-open", filepath], check=True)
            return f"and opened it for you."
        except Exception as e:
            return f"but I couldn't open it automatically. Error: {e}"

    def _handle_file_creation(self, action, params):
        filename = params.get("filename"); content_topic = params.get("content", ""); is_topic = params.get("is_topic", False)
        if not filename: return "I need a filename to create a file."
        filepath = self.workspace_dir / filename
        try:
            if action == "create_excel":
                if not is_topic: return "Please provide a topic for the Excel data."
                json_data_str = self._generate_content_with_ai(content_topic, output_format='json')
                match = re.search(r'\{.*\}', json_data_str, re.DOTALL)
                if not match: raise json.JSONDecodeError("No JSON object found in AI response.", json_data_str, 0)
                data = json.loads(match.group(0))
                workbook = Workbook(); sheet = workbook.active
                sheet.append(data['headers'])
                for row in data['rows']: sheet.append(row)
                workbook.save(filepath)
            elif action == "create_python":
                final_content = self._generate_content_with_ai(content_topic, output_format='code') if is_topic else content_topic
                filepath.write_text(final_content, encoding='utf-8')
            elif action in ["create_word", "create_pdf", "create_text"]:
                final_content = self._generate_content_with_ai(content_topic, output_format='text') if is_topic else content_topic
                if action == "create_text": filepath.write_text(final_content, encoding='utf-8')
                elif action == "create_word": self._create_formatted_word_doc(filepath, final_content)
                elif action == "create_pdf": self._create_formatted_pdf(filepath, final_content)
            else:
                return f"Unknown file creation action: {action}"
            open_status = self._open_file(filepath)
            return f"✅ Successfully created '{filename}' {open_status}"
        except json.JSONDecodeError as e:
            return f"❌ The AI returned invalid data for the Excel sheet. Error: {e}"
        except Exception as e:
            return f"❌ Failed to create file '{filename}'. Error: {e}"

    def _create_formatted_word_doc(self, filepath, markdown_content):
        doc = docx.Document()
        for line in markdown_content.split('\n'):
            line = line.strip()
            if line.startswith('## '): doc.add_heading(line[3:], level=2)
            elif line.startswith('# '): doc.add_heading(line[2:], level=1)
            elif line.startswith('* '): doc.add_paragraph(line[2:], style='List Bullet')
            elif line:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'): p.add_run(part[2:-2]).bold = True
                    elif part.startswith('*') and part.endswith('*'): p.add_run(part[1:-1]).italic = True
                    else: p.add_run(part)
        doc.save(filepath)

    def _create_formatted_pdf(self, filepath, markdown_content):
        pdf = FPDF()
        pdf.add_page(); pdf.set_auto_page_break(auto=True, margin=15)
        try:
            pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
            pdf.set_font('DejaVu', '', 12)
        except RuntimeError:
            print("⚠️ WARNING: DejaVuSans.ttf not found. PDF may have character errors.")
            pdf.set_font('Arial', '', 12)
        for line in markdown_content.split('\n'):
            line = line.strip()
            if line.startswith('# '): pdf.set_font(size=18); pdf.multi_cell(0, 12, line[2:], new_x="LMARGIN", new_y="NEXT")
            elif line.startswith('## '): pdf.set_font(size=14); pdf.multi_cell(0, 10, line[3:], new_x="LMARGIN", new_y="NEXT")
            elif line.startswith('* '): pdf.set_font(size=12); pdf.multi_cell(0, 8, f'• {line[2:]}', new_x="LMARGIN", new_y="NEXT", markdown=True)
            elif line: pdf.set_font(size=12); pdf.multi_cell(0, 8, line, new_x="LMARGIN", new_y="NEXT", markdown=True)
            else: pdf.ln(5)
        pdf.output(filepath)

    def _handle_system_control(self, action, params):
        app_aliases = {"visual studio code": "code", "vscode": "code", "word": "winword", "excel": "excel"}
        if action == "open_application":
            app_name = params.get("application", "").lower()
            if not app_name: return "You need to specify an application to open."
            command_to_run = app_aliases.get(app_name, app_name)
            try:
                print(f"⚙️  Attempting to run '{command_to_run}'...")
                if self.system == "windows": os.system(f'start {command_to_run}')
                else: subprocess.Popen([command_to_run])
                return f"{app_name.capitalize()} opened."
            except Exception as e: return f"Sorry, I couldn't open {app_name}. Error: {e}"
        return "System command executed."

    def _handle_web_Browse(self, action, params):
        query = params.get("search_query", "");
        if query: webbrowser.open(f"https://www.google.com/search?q={query.replace(' ', '+')}"); return f"Searching the web for: {query}"
        else: webbrowser.open("https://www.google.com"); return "I wasn't given a query, so I've opened Google.com."
            
    def _handle_file_management(self, action, params):
        if action == "list_files": files = list(self.workspace_dir.glob("*")); return f"Files in workspace:\n" + "\n".join([f"  📄 {f.name}" for f in files]) if files else "Your workspace is empty."
        elif action == "find_file":
            filename = params.get("filename");
            if not filename: return "What file do you want to find?"
            print(f"🔎 Searching for '{filename}'..."); results = list(Path.home().rglob(filename))
            if not results: return f"I couldn't find '{filename}'."
            return f"I found:\n" + "\n".join([f"  -> {p}" for p in results[:10]])
        elif action == "delete_file":
            filename = params.get("filename");
            if not filename: return "What file do you want to delete?"
            results = list(Path.home().rglob(filename));
            if not results: return f"I couldn't find '{filename}'."
            if len(results) > 1: return "I found multiple files. Please be more specific."
            self.pending_confirmation = {"action": "confirm_delete", "file_path": results[0]}
            return f"I found: {results[0]}\n\n⚠️ Are you sure you want to permanently delete it? (yes/no)"
        return "Unknown file management action."

    def _confirm_delete_action(self, details):
        try: Path(details["file_path"]).unlink(); return f"✅ Successfully deleted."
        except Exception as e: return f"❌ Failed to delete file. Error: {e}"

    def _handle_conversation(self, params):
        content = params.get("content", "").lower()
        if "how are you" in content:
            return "I am functioning at optimal capacity. Ready for your command."
        return "I am JARVIS, your AI assistant. How can I help you create, manage, or find information?"
    
    def run(self):
        print("\nJARVIS AI Assistant is active. How may I assist you?")
        while True:
            try:
                print("\n" + "="*50)
                if self.pending_confirmation:
                    prompt_msg = self.pending_confirmation.get('prompt', "Please confirm (yes/no): ")
                    print(f"🤖 JARVIS: {prompt_msg}")
                    user_response = input("💬 Your confirmation (yes/no): ").strip().lower()
                    if user_response == 'yes':
                        if self.pending_confirmation["action"] == "confirm_delete": print(f"🤖 JARVIS: {self._confirm_delete_action(self.pending_confirmation)}")
                    else: print("🤖 JARVIS: OK, action cancelled.")
                    self.pending_confirmation = None
                    continue
                command = input("💬 Type your command (or 'exit' to quit): ").strip()
                if not command: continue
                if command.lower() in ['exit', 'quit', 'shutdown']: break
                analysis = self.analyze_command_with_ai(command)
                print(f"📊 Intent: {analysis.get('intent', 'N/A')} | Action: {analysis.get('action', 'N/A')}")
                result = self.execute_command(analysis)
                if self.pending_confirmation: self.pending_confirmation['prompt'] = result
                else: print(f"🤖 JARVIS: {result}")
            except KeyboardInterrupt: break
            except Exception as e: print(f"❌ An unexpected error occurred: {e}")
        print("\nShutting down JARVIS. Goodbye!")

def main():
    print("🚀 Initializing JARVIS AI Assistant...")
    print("📋 Requirements: pip install google-generativeai python-docx openpyxl fpdf2")
    api_key = "" 
    if not api_key or api_key == "":
        print("\n🔑 WARNING: API key not set! Please paste your Google Gemini API key to enable AI features.\n")
    try:
        jarvis = JarvisAI(api_key=api_key)
        jarvis.run()
    except Exception as e:
        print(f"❌ An error occurred while starting JARVIS: {e}")

if __name__ == "__main__":
    main()