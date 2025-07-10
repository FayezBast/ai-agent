# file_manager.py
import re
import json
from pathlib import Path
from typing import List, Dict

# Third-party imports with availability checks
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import Workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

from config import Config
from security import SecurityValidator
from logger import log_info, log_error, log_warning

class FileManager:
    """Handles all file operations within the designated workspace."""

    def __init__(self):
        self.workspace_dir = Config.WORKSPACE_DIR
        self.workspace_dir.mkdir(exist_ok=True)
        log_info(f"File manager initialized. Workspace: {self.workspace_dir}")

    def create_file(self, filename: str, content: str, file_type: str) -> str:
        """
        Creates a file with the given content and type after validation.
        """
        try:
            validated_filename = SecurityValidator.validate_filename(filename)
            filepath = self.workspace_dir / validated_filename
            
            # Ensure the path is still within the workspace after combining
            SecurityValidator.validate_path(filepath)

            sanitized_content = SecurityValidator.sanitize_content(content)

            if file_type == 'docx' and DOCX_AVAILABLE:
                self._create_word_doc(filepath, sanitized_content)
            elif file_type == 'xlsx' and EXCEL_AVAILABLE:
                self._create_excel_doc(filepath, sanitized_content)
            elif file_type == 'pdf' and PDF_AVAILABLE:
                self._create_pdf_doc(filepath, sanitized_content)
            else: # Default to text or specified type like .py, .txt
                self._create_text_based_file(filepath, sanitized_content)
            
            log_info(f"Successfully created file: {filepath}")
            return f"Successfully created '{validated_filename}' in your workspace."
        except ValueError as ve:
            log_error(f"File creation validation error: {ve}")
            return f"Error: {ve}"
        except Exception as e:
            log_error(f"File creation failed for {filename}: {e}")
            return f"An unexpected error occurred while creating the file: {e}"

    def find_files(self, query: str) -> List[Path]:
        """Finds files in the workspace matching a query."""
        log_info(f"Searching for files with query: '{query}'")
        try:
            # Use rglob for recursive search within the workspace
            return list(self.workspace_dir.rglob(f"*{query}*"))
        except Exception as e:
            log_error(f"Error during file search: {e}")
            return []

    def delete_file(self, file_path: Path) -> str:
        """Deletes a file after validating its path."""
        try:
            validated_path = SecurityValidator.validate_path(file_path)
            if validated_path.exists() and validated_path.is_file():
                validated_path.unlink()
                log_info(f"Deleted file: {validated_path}")
                return f"File '{validated_path.name}' has been deleted."
            else:
                log_warning(f"Attempted to delete non-existent file: {file_path}")
                return "File not found."
        except ValueError as ve:
            log_error(f"File deletion validation error: {ve}")
            return f"Error: {ve}"
        except Exception as e:
            log_error(f"Failed to delete file {file_path}: {e}")
            return f"An error occurred during file deletion: {e}"

    def list_workspace_files(self) -> List[Dict[str, str]]:
        """Lists all files in the root of the workspace with their details."""
        files_details = []
        try:
            for item in self.workspace_dir.iterdir():
                if item.is_file():
                    stat = item.stat()
                    files_details.append({
                        "name": item.name,
                        "size": f"{stat.st_size / 1024:.2f} KB",
                        "modified": str(item.stat().st_mtime)
                    })
            return files_details
        except Exception as e:
            log_error(f"Failed to list workspace files: {e}")
            return []

    def _create_text_based_file(self, filepath: Path, content: str):
        filepath.write_text(content, encoding='utf-8')

    def _create_word_doc(self, filepath: Path, content: str):
        doc = docx.Document()
        doc.add_heading(filepath.stem, level=1)
        # Simple parsing for markdown-like text
        for line in content.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
        doc.save(filepath)

    def _create_excel_doc(self, filepath: Path, content: str):
        workbook = Workbook()
        sheet = workbook.active
        try:
            # Assume content is JSON with 'headers' and 'rows'
            data = json.loads(content)
            sheet.title = filepath.stem[:30]
            sheet.append(data['headers'])
            for row in data['rows']:
                sheet.append(row)
        except (json.JSONDecodeError, KeyError):
            # Fallback for non-JSON content
            sheet.title = "Content"
            sheet['A1'] = "Generated Content"
            sheet['A2'] = content
        workbook.save(filepath)

    def _create_pdf_doc(self, filepath: Path, content: str):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, content)
        pdf.output(str(filepath))